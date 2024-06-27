##################################
#Importaciones
##################################
import streamlit as st
import pandas as pd
import numpy as np
import spacy
import textdescriptives
from docx import Document
from wordcloud import WordCloud
nlp = spacy.load("es_core_news_lg")
nlp.add_pipe("textdescriptives/all")
nlp.disable_pipes('ner')
import re
from wordcloud import WordCloud
from collections import Counter
import matplotlib.pyplot as plt
from pysentimiento import create_analyzer
analyzer = create_analyzer(task="sentiment", lang="es")
emotion_analyzer = create_analyzer(task="emotion", lang="es")
from spacy.tokens import Doc
import csv
import io
import time
from io import BytesIO
import openpyxl

##################################
#Definición de funciones
##################################

#Leer un documento de word
def leer_docx(nombre_archivo):
    texto = ""
    doc = Document(nombre_archivo)
    for paragraph in doc.paragraphs:
        texto += paragraph.text + "\n"
    return texto


#Para agregar stopwords y lematizar
def stopYlematizar(frase):
    frase = frase.lower()
    doc = nlp(frase)
    lista_limpia = [token.lemma_ for token in doc if token.is_alpha]
    frase = " ".join(lista_limpia)
    doc = nlp(frase)
    lista_limpia = [token.text for token in doc if not token.is_space and not token.is_stop and not token.is_punct]
    return lista_limpia

#Devuelve la temática
def semantica(frase, tipo):
    frase = frase.lower()
    doc = nlp(frase)
    lista = [token.text for token in doc if token.pos_ == tipo]
    return lista

#Devuelve las entidades
def entidades(frase):
    frase = frase.lower()
    doc = nlp(frase)
    lista = [ent.text for ent in doc.ents]
    return lista

#Lee un txt
def leer_txt(nombre_archivo):
    # Lista de codificaciones comunes a probar
    codificaciones = ['utf-8']

    for codificacion in codificaciones:
            with open(".\\Libros\\100 años de soledad.txt", 'r', encoding=codificacion) as archivo:
                # Lee el contenido del archivo
                texto = archivo.read()
    return texto

#Devuelve el coeficiente de facilidad lectora
def coeficiente_facilidad_lectora(doc):
    p = doc._.syllables['syllables_per_token_mean']
    f = doc._.sentence_length['sentence_length_mean']
    return 206.84 - 60 * p - 1.02 * f

def nubePalabras(frase):
    w = WordCloud()
    w.generate(frase)

    fig, ax = plt.subplots(1, 1, figsize=(20, 7), dpi=100)
    ax.imshow(w, interpolation='bilinear')
    ax.axis("off")  # No mostrar los ejes

    # Mostrar la figura en Streamlit
    st.pyplot(fig)

def frecuencias(lista):
    df = pd.DataFrame(lista)
    conteo_frecuencias = df.value_counts()
    frecuencias = conteo_frecuencias.to_frame()
    frecuencias.reset_index(inplace = True)
    frecuencias.columns = ['token', 'conteo']
    frecuencias = frecuencias.sort_values(by='conteo', ascending=False)
    return frecuencias

#Divide por capítulos
def dividir_x_capitulos(indice, libro_x_frase):
    lista_capitulos = []
    lista_capitulo = []
    if len(indice) > 0:
        patron = re.compile("|".join(map(re.escape, indice)))
        for parrafo in libro_x_frase:
            if bool(patron.search(parrafo)) == True:
                if len(lista_capitulo) > 0:
                    lista_capitulos.append(lista_capitulo.copy())
                    lista_capitulo.clear()
            else:
                lista_capitulo.append(parrafo)

        if len(lista_capitulo) > 0:
            lista_capitulos.append(lista_capitulo.copy())
            lista_capitulo.clear()
    else:
        lista_capitulos.append(libro_x_frase)
    return lista_capitulos

# Función para leer el archivo Excel y obtener la lista de capítulos
def obtener_capitulos(uploaded_file):
    # Leer el archivo Excel
    df = pd.read_excel(uploaded_file, header=None)
    # Suponiendo que los capítulos están en la primera columna
    capitulos = df.iloc[:, 0].tolist()
    return capitulos

def generar_csv(df):
    csv_buffer = io.StringIO()
    df.to_csv(csv_buffer, index=False)
    csv_data = csv_buffer.getvalue()
    return csv_data

# Estado inicial
if 'graficas_generadas' not in st.session_state:
    st.session_state.graficas_generadas = True

##################################
# Funcionalidad
##################################
st.title('Análisis estadístico de libros en español')

st.write("Hola, esta es una página gratuita para mostrar el análisis estadístico de cualquier libro que esté en español.")

st.write("El libro que puedo recibir puede estar en formato .txt o .docx, te recomiendo que dejes solo el texto que desees analizar "
         "y excluir cosas como la editorial, el título, el prólogo, agradecimientos, etc. También, en casos de Word, te recomiendo borrar "
         "el encabezado y pie de página, ya que Python lo detecta como texto y puede alterar los datos.")

st.write("Ahora que ya quedó claro, en el siguiente campo sube el archivo del libro:")

# Subir el archivo
uploaded_file = st.file_uploader("Elige un archivo", type=["txt", "docx"])

st.write("Tienes la posibilidad de analizar capítulos individuales, para ello, sube en el siguiente campo una lista en .xlsx donde cada fila sea un capítulo."
         " Esto es totalmente opcional y puede resultar tardado dependiendo el largo del capítulo, así que ten paciencia.")

capitulos_file = st.file_uploader("Elige un archivo", type=["xlsx"])

if uploaded_file is not None and capitulos_file is not None:
    if st.button("Analizar libro por capítulos"):
        if uploaded_file.type == "text/plain":
            # Decodificar contenido del archivo de texto
            content = uploaded_file.read().decode("utf-8")

        # Leer archivo de Word
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Leer el documento Word usando la librería python-docx
            doc = Document(uploaded_file)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            content = "\n".join(full_text)
        else:
            st.write("xxErrorxx: El archivo no tiene la extensión permitida: .txt o docx")

        #Si pasa este if, podemos continuar
        if len(content) > 0:
            st.session_state.graficas_generadas = True
            if st.session_state.graficas_generadas:   
                st.text_area("Contenido del libro", content, height=300)
                #Vamos a crear la estadística del libro completo
                libro_x_frase = content.split("\n")
                libro_x_frase = list(filter(None, libro_x_frase))

                if capitulos_file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                    # Decodificar contenido del archivo de texto
                    indice = obtener_capitulos(capitulos_file)
                else:
                    st.write("xxErrorxx: El archivo de índice no tiene la extensión permitida: .xlsx")

                if len(indice) > 0:
                    #Comenzamos el análisis
                    lista_capitulos = dividir_x_capitulos(indice, libro_x_frase)

                    libro_dic = {
                        'Capitulo' : [],
                        'Frases' : [],
                        'palabras' : [],
                        'sustantivos' : [],
                        'NEU' : [],
                        'POS' : [],
                        'NEG' : [],
                        'Felicidad' : [],
                        'Enojo' : [],
                        'Disgusto' : [],
                        'Tristeza' : [],
                        'Sorpresa' : [],
                        'Miedo' : [],
                        'Otros' : [],
                        'n_tokens' : [],
                        'n_unique_tokens' : [],
                        'proportion_unique_tokens' : [],
                        'n_characters' : [],
                        'n_sentences' : [],
                        'sentence_length_mean' : [],
                        'sentence_length_median' : [],
                        'syllables_per_token_mean' : [],
                        'syllables_per_token_median' : [],
                        'token_length_mean' : [],
                        'token_length_median' : [],
                        'coef_facilidad_lect' : []
                    }

                    capitulos_completos = []
                    count = 0
                    for capitulo in lista_capitulos:
                        capitulo_completo = " ".join(capitulo)
                        capitulos_completos.append(capitulo_completo)

                    st.write(f'Capítulos encontrados en índice:{len(indice)}')
                    st.write(f'Capítulos encontrados en libro:{len(capitulos_completos)}')

                    if len(indice) != len(capitulos_completos):
                        st.write(f'Analize la división de capítulos para realizar las modificaciones necesarias')
                        st.write(f'Índice')
                        st.write(indice)
                        st.write(f'División de capítulos')
                        st.write(capitulos_completos)
                    else:
                        for capitulo in capitulos_completos:
                            capitulo_completo = " ".join(capitulo)
                            st.write(f'Analizando capítulo: {indice[count]}...')
                            sentimientos = {
                                'NEU' : [],
                                'POS' : [],
                                'NEG' : []
                            }
                            emociones = {
                                'Felicidad' : [],
                                'Enojo' : [],
                                'Disgusto' : [],
                                'Tristeza' : [],
                                'Sorpresa' : [],
                                'Miedo' : [],
                                'Otros' : []
                            }
                            doc = nlp(capitulo)
                            libro_dic["Capitulo"].append(indice[count])
                            count += 1
                            libro_dic["Frases"].append(len(list(doc.sents)))
                            libro_dic["palabras"].append(", ".join(map(str, Counter(stopYlematizar(capitulo)).most_common(3))).replace("(", "").replace(")", "").replace("',", "':").replace("'", ""))
                            libro_dic["sustantivos"].append(", ".join(map(str, Counter(semantica(capitulo, "NOUN")).most_common(3))).replace("(", "").replace(")", "").replace("',", "':").replace("'", ""))
                            for oracion in list(doc.sents):
                                a = analyzer.predict(oracion.text)
                                sentimientos["NEU"].append(float(a.probas["NEU"]))
                                sentimientos["POS"].append(float(a.probas["POS"]))
                                sentimientos["NEG"].append(float(a.probas["NEG"]))
                                e = emotion_analyzer.predict(oracion.text)
                                emociones["Felicidad"].append(float(e.probas["joy"]))
                                emociones["Enojo"].append(float(e.probas["anger"]))
                                emociones["Disgusto"].append(float(e.probas["disgust"]))
                                emociones["Tristeza"].append(float(e.probas["sadness"]))
                                emociones["Sorpresa"].append(float(e.probas["surprise"]))
                                emociones["Miedo"].append(float(e.probas["fear"]))
                                emociones["Otros"].append(float(e.probas["others"]))
                            df_a = pd.DataFrame(sentimientos)
                            df_e = pd.DataFrame(emociones)
                            libro_dic["NEU"].append(df_a["NEU"].mean())
                            libro_dic["POS"].append(df_a["POS"].mean())
                            libro_dic["NEG"].append(df_a["NEG"].mean())
                            libro_dic["Felicidad"].append(df_e["Felicidad"].mean())
                            libro_dic["Enojo"].append(df_e["Enojo"].mean())
                            libro_dic["Disgusto"].append(df_e["Disgusto"].mean())
                            libro_dic["Tristeza"].append(df_e["Tristeza"].mean())
                            libro_dic["Sorpresa"].append(df_e["Sorpresa"].mean())
                            libro_dic["Miedo"].append(df_e["Miedo"].mean())
                            libro_dic["Otros"].append(df_e["Otros"].mean())
                            libro_dic["n_tokens"].append(doc._.counts['n_tokens'])
                            libro_dic["n_unique_tokens"].append(doc._.counts['n_unique_tokens'])
                            libro_dic["proportion_unique_tokens"].append(doc._.counts['proportion_unique_tokens'])
                            libro_dic["n_characters"].append(doc._.counts['n_characters'])
                            libro_dic["n_sentences"].append(doc._.counts['n_sentences'])
                            libro_dic["sentence_length_mean"].append(doc._.sentence_length['sentence_length_mean'])
                            libro_dic["sentence_length_median"].append(doc._.sentence_length['sentence_length_median'])
                            libro_dic["syllables_per_token_mean"].append(doc._.syllables['syllables_per_token_mean'])
                            libro_dic["syllables_per_token_median"].append(doc._.syllables['syllables_per_token_median'])
                            libro_dic["token_length_mean"].append(doc._.token_length['token_length_mean'])
                            libro_dic["token_length_median"].append(doc._.token_length['token_length_median'])
                            libro_dic["coef_facilidad_lect"].append(coeficiente_facilidad_lectora(doc))

                        df = pd.DataFrame(libro_dic)
                        df_libro_x_cap = df.copy()

                        columnas = ['Capítulo', 'Frases', 'Palabras más comunes', 'Sustantivos más comunes',
                                    'Sentimiento Neutral', 'Sentimiento Positivo', 'Sentimiento Negativo', 'Emoción felicidad',
                                    'Emoción enojo', 'Emoción disgusto', 'Emoción tristeza', 'Emoción sorpresa', 'Emoción miedo', 'Emoción otros',
                                    'Número de palabras', 'Número de palabras únicas', 'Proporción de palabras únicas', 'Número de letras',
                                    'Número de enunciados', 'Promedio de longitud enunciados', 'Mediana de longitud enunciados', 'Promedio de sílabas por palabra',
                                    'Mediana de sílabas por palabra', 'Promedio de longitud de palabras', 'Mediana de logitud de palabras', 'Coeficiente de facilidad de lectura']
                        
                        df.columns = columnas

                        #st.write("Análisis por Capítulo:")
                        #st.dataframe(df)

                        # Colores y estilos de línea accesibles para daltónicos
                        styles = {
                            "Felicidad": {"color": "gold", "linestyle": "-"},  
                            "Enojo": {"color": "red", "linestyle": "--"},     
                            "Disgusto": {"color": "green", "linestyle": "-."},  
                            "Tristeza": {"color": "mediumblue", "linestyle": ":"},   
                            "Sorpresa": {"color": "mediumorchid", "linestyle": "-"},   
                            "Miedo": {"color": "steelblue", "linestyle": "--"}      
                        }

                        #st.write("Gráfica de Análisis de Emociones por Capítulo")

                        fig1, ax1 = plt.subplots(figsize=(20, 6))
                        ax1.plot(df_libro_x_cap.index + 1, df_libro_x_cap["Felicidad"], label="Felicidad", **styles["Felicidad"])
                        ax1.plot(df_libro_x_cap.index + 1, df_libro_x_cap["Enojo"], label="Enojo", **styles["Enojo"])
                        ax1.plot(df_libro_x_cap.index + 1, df_libro_x_cap["Disgusto"], label="Disgusto", **styles["Disgusto"])
                        ax1.plot(df_libro_x_cap.index + 1, df_libro_x_cap["Tristeza"], label="Tristeza", **styles["Tristeza"])
                        ax1.plot(df_libro_x_cap.index + 1, df_libro_x_cap["Sorpresa"], label="Sorpresa", **styles["Sorpresa"])
                        ax1.plot(df_libro_x_cap.index + 1, df_libro_x_cap["Miedo"], label="Miedo", **styles["Miedo"])

                        ax1.set_xlabel('Capítulos')
                        ax1.set_ylabel('Razón de Emoción')
                        ax1.set_xticks(np.arange(1, len(df_libro_x_cap) + 1, step=1))
                        ax1.legend()
                        ax1.grid(True)

                        #st.pyplot(fig1)

                        styles = {
                            "POS": {"color": "gold", "linestyle": "-"},  
                            "NEG": {"color": "dimgrey", "linestyle": "--"}   
                        }

                        #st.write("Gráfica de Análisis de Sentimientos por Capítulo")

                        fig2, ax2 = plt.subplots(figsize=(20, 6))
                        ax2.plot(df_libro_x_cap.index + 1, df_libro_x_cap["POS"], label="Positivo", **styles["POS"])
                        ax2.plot(df_libro_x_cap.index + 1, df_libro_x_cap["NEG"], label="Negativo", **styles["NEG"])

                        ax2.set_xlabel('Capítulos')
                        ax2.set_ylabel('Razón de Sentimiento')
                        ax2.set_xticks(np.arange(1, len(df_libro_x_cap) + 1, step=1))
                        ax2.legend()
                        ax2.grid(True)

                        #st.pyplot(fig2) 

                        st.write("Generando reporte...")

                        output = BytesIO()
                        with pd.ExcelWriter(output, engine='openpyxl') as writer:
                            df.to_excel(writer, index=True, sheet_name='Capítulos')
                            workbook = writer.book
                            worksheet = workbook.create_sheet(title='Sentimientos')
                            imgdata = BytesIO()
                            fig2.savefig(imgdata, format='png')
                            imgdata.seek(0)
                            image = openpyxl.drawing.image.Image(imgdata)
                            worksheet.add_image(image, 'B2')

                            worksheet = workbook.create_sheet(title='Emociones')
                            imgdata = BytesIO()
                            fig1.savefig(imgdata, format='png')
                            imgdata.seek(0)
                            image = openpyxl.drawing.image.Image(imgdata)
                            worksheet.add_image(image, 'B2')
                        output.seek(0)

                        # Botón para descargar el archivo Excel
                        st.download_button(
                            label="Descargar Excel",
                            data=output,
                            file_name='análisis de capítulos.xlsx',
                            mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                        )               
elif uploaded_file is not None:
    if st.button("Analizar libro completo"):
        if uploaded_file.type == "text/plain":
            # Decodificar contenido del archivo de texto
            content = uploaded_file.read().decode("utf-8")

        # Leer archivo de Word
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Leer el documento Word usando la librería python-docx
            doc = Document(uploaded_file)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            content = "\n".join(full_text)
        else:
            st.write("xxErrorxx: El archivo no tiene la extensión permitida: .txt o docx")

        #Si pasa este if, podemos continuar
        if len(content) > 0:
            st.session_state.graficas_generadas = True
            if st.session_state.graficas_generadas:
                st.text_area("Contenido del libro", content, height=300)

                #Vamos a crear la estadística del libro completo
                libro_x_frase = content.split("\n")
                libro_x_frase = list(filter(None, libro_x_frase))

                libro = " ".join(libro_x_frase)

                st.write("Analizando el libro... esto puede tardar.")
                # Obtenemos la lista de palabras con lematización y sin stopwords
                libro_lista_lemm = stopYlematizar(libro)
                
                st.write("Nube de palabras")
                # Creamos el wordcloud con todas las palabras con lematización y stopwords
                nubePalabras(" ".join(libro_lista_lemm))

                frecuencias = frecuencias(libro_lista_lemm)

                st.write("Gráfica de frecuencias")

                # Crear la figura y el gráfico de barras
                fig, ax = plt.subplots(figsize=(10, 6))
                ax.bar(frecuencias.iloc[:20].token, frecuencias.iloc[:20].conteo)
                ax.set_xlabel('Palabra')
                ax.set_ylabel('Frecuencia')
                ax.set_xticks(range(len(frecuencias.iloc[:20].token)))
                ax.set_xticklabels(frecuencias.iloc[:20].token, rotation=45)
                st.pyplot(fig)